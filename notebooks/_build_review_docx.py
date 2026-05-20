"""Deep-dive review of the forecasting + churn methodology.

Separates methodology problems (things we built that have flaws) from
data problems (things missing from the source data). For each issue:
what it is in plain language, what it costs us, and the fix.

Output: notebooks/output/Pryzm_Methodology_Review.docx
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = Path(__file__).resolve().parent / "output"
REPORT = OUT_DIR / "Pryzm_Methodology_Review.docx"

# --- colors ----------------------------------------------------------------

INK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x60, 0x6B, 0x7A)
ROSE = RGBColor(0x3E, 0x5D, 0x80)
RED = RGBColor(0xB3, 0x26, 0x1A)
AMBER = RGBColor(0xB8, 0x70, 0x10)
GREEN = RGBColor(0x10, 0x80, 0x40)


def heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = ROSE if level <= 1 else INK
        r.font.name = "Calibri"


def para(doc: Document, text: str, *, bold: bool = False, italic: bool = False,
         color: RGBColor | None = None, size: int = 11) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def issue(doc: Document, *, number: str, title: str, severity: str,
          plain: str, impact: str, fix: str, tags: list[str]) -> None:
    """Render one structured issue card."""
    sev_color = {"HIGH": RED, "MEDIUM": AMBER, "LOW": GREEN}.get(severity, INK)

    # Title line
    p = doc.add_paragraph()
    n = p.add_run(f"{number}.  ")
    n.bold = True
    n.font.size = Pt(13)
    n.font.color.rgb = INK
    t = p.add_run(title)
    t.bold = True
    t.font.size = Pt(13)
    t.font.color.rgb = INK

    # Severity + tags row
    p = doc.add_paragraph()
    s = p.add_run(f"Severity: {severity}")
    s.bold = True
    s.font.color.rgb = sev_color
    s.font.size = Pt(10)
    if tags:
        tag_text = "    Tags: " + " · ".join(tags)
        t = p.add_run(tag_text)
        t.font.color.rgb = MUTED
        t.font.size = Pt(10)
        t.italic = True

    for label, text, color in [
        ("What it is", plain, INK),
        ("Why it matters", impact, INK),
        ("Fix", fix, GREEN),
    ]:
        p = doc.add_paragraph()
        l = p.add_run(f"{label}: ")
        l.bold = True
        l.font.size = Pt(11)
        l.font.color.rgb = color
        t = p.add_run(text)
        t.font.size = Pt(11)

    doc.add_paragraph()  # spacer


# ---------------------------------------------------------------------------

METHODOLOGY_ISSUES = [
    dict(
        number="M1",
        title="Quote conversion model has feature leakage — AUC is inflated",
        severity="HIGH",
        tags=["leakage", "validation"],
        plain="Each quote gets a feature called 'cust_winrate_hist' = the customer's overall win-rate across ALL their quotes. We compute this once from the full training set and apply it to every row, including the row we're trying to predict. So a customer's historical win-rate feature contains the outcome of the very quote we're scoring.",
        impact="The reported AUC of 0.84 likely overstates real out-of-sample performance. On a true holdout the AUC could drop to 0.65-0.75. The pipeline_forecast.csv expected-revenue numbers inherit this optimism. We are quoting a precision the model doesn't have.",
        fix="Replace single global customer_history with a per-row leave-one-out computation: for each quote, win-rate uses ONLY quotes from the same customer that closed strictly before the quote's date. Costs ~2x training time. Re-validate AUC; expect a drop. Document the realistic AUC in the validation report.",
    ),
    dict(
        number="M2",
        title="Quantile bands assume Normal residuals — they don't match real B2B revenue",
        severity="HIGH",
        tags=["uncertainty", "calibration"],
        plain="Every p80 / p95 band is derived as 'forecast ± k × residual_std' where k comes from a standard Normal distribution (1.28 for p80, ~1.96 for p95). This assumes errors are symmetric and bell-curved. Industrial B2B revenue / quantity series are right-skewed (occasional huge orders) and heavy-tailed (long quiet stretches punctuated by big shipments).",
        impact="The 80% bands are systematically too narrow on the upside, too wide on the downside, for monetary metrics. When a forecast says 'p80 = €100k ± €30k' the actual realisation routinely sits outside the band. Plans built around p80 will be wrong more often than 1-in-5.",
        fix="Switch to bootstrap residual quantiles: compute empirical 10th/90th percentiles of in-sample errors directly, without the Normal assumption. For log-transformed series, do quantiles in log-space then expm1 (already partially done but not for the band derivation). Long term, fit quantile regression directly (statsmodels has QuantReg, or use scikit GradientBoostingRegressor(loss='quantile').",
    ),
    dict(
        number="M3",
        title="Model selection uses backtest MAPE — but we then report that same MAPE",
        severity="MEDIUM",
        tags=["selection bias", "validation"],
        plain="For each series we fit 2-4 candidate models, pick the one with the lowest walk-forward MAPE, then report that MAPE as the model's accuracy. The MAPE we report is the WINNER's, which is by construction the best of several tries on the same data.",
        impact="Reported MAPEs are biased optimistic by 1-3 percentage points (more for short series with few backtest points). We may claim the gate is met when, on truly fresh data, it isn't. Particularly relevant for BKAIZ at 10.5% MAPE — the true MAPE could be 12-13%.",
        fix="Nested split: use an inner holdout to select the model, then an outer holdout (held-out from selection) to estimate reported accuracy. Or use cross-validated AIC for selection and reserve the holdout purely for measurement. Implement once and apply uniformly.",
    ),
    dict(
        number="M4",
        title="Per-customer / per-SKU monthly forecasts use the wrong model family",
        severity="HIGH",
        tags=["model fit", "informational tier"],
        plain="Bursty B2B per-account series have many zero / near-zero months interrupted by large orders. We fit EMA / SES / Linear-trend / ETS — all built for continuous, smoothly-varying series. None of these handle intermittent demand. The result is the 55–80% median MAPE we documented as 'inherently bursty'.",
        impact="The 'informational tier' label hides the real problem: we're using the wrong tool. With the right tool the per-customer / per-SKU forecasts could shift from informational to usable, which would lift the at-risk-revenue calculation from rough magnitude to actual euros.",
        fix="Add Croston's method and TSB (Teunter-Syntetos-Babai) as candidate models for series with >30% zero months. These are the industry standards for intermittent demand. statsforecast (open source) has tuned implementations. Should drop the customer-level WAPE from ~55% to ~30%.",
    ),
    dict(
        number="M5",
        title="Hierarchical reconciliation is naive bottom-up scaling, not MinT",
        severity="MEDIUM",
        tags=["coherence", "statistical efficiency"],
        plain="To make customer / SKU forecasts add up to commodity-group forecasts, we compute the ratio of (sum of children) to (parent) each month and multiply every child's p50/p80/p95 by that ratio. This is the simplest coherence trick. The reconciliation residuals report shows pre-scaling residuals of 100%+ for some groups — that's a lot of forced adjustment.",
        impact="Information is being thrown away. A customer with a strong local trend can get its p50 doubled by the reconciliation because the parent says total revenue should be higher. The resulting forecast is mathematically consistent but loses the per-customer signal. Bands also get blindly stretched.",
        fix="Replace with MinT (Minimum Trace) reconciliation: weights residuals by their inverse covariance so high-confidence children pull more weight. Implementation: hierarchicalforecast package (statsforecast ecosystem) or roll our own — it's a single matrix operation.",
    ),
    dict(
        number="M6",
        title="Quarterly forecasts assume monthly errors are independent",
        severity="MEDIUM",
        tags=["uncertainty", "MC"],
        plain="We derive quarterly bands by drawing 1,000 sample paths of monthly values (each month as an independent Normal draw around its p50) and summing into quarters. But monthly errors in the same series are autocorrelated — if January is high, February tends to be high too.",
        impact="Quarterly bands look narrower than they should. The variance of a sum of correlated variables is larger than the sum of variances; we're using the lower bound. Quarterly p80 bands are systematically too tight — a customer whose Q1 lands outside the p80 is over-represented vs the 1-in-5 promise.",
        fix="Either fit a model that directly produces quarterly forecasts (ETS on quarterly-aggregated series), or sample monthly residuals from a fitted multivariate Normal (or use a block bootstrap that preserves autocorrelation). Easier route: just fit ETS to quarterly aggregates and use those directly for the quarterly view.",
    ),
    dict(
        number="M7",
        title="SARIMAX-with-exog never wins — but we don't know if that's the methodology or the data",
        severity="MEDIUM",
        tags=["macro regressors", "investigation"],
        plain="We built SARIMAX-with-exog wired to the correlation_map, expecting BKAIZ (the laggard) to drop below 8% MAPE. Zero series got upgraded. The incumbent SES models always won the backtest.",
        impact="We invested in the macro-regressor pipeline and got no measurable lift. Either (a) the macro indicators don't actually have signal at this aggregation level, (b) the lag detection in correlation_map is overfit, (c) the (1,1,1) ARIMA order is wrong, or (d) the exog forward-projection (random-walk-with-drift) discards the lead too aggressively. We don't know which.",
        fix="Run a focused diagnostic: pick BKAIZ, fit SARIMAX with each leading indicator one at a time, vary the order, log every backtest. Either we find the right combination and BKAIZ drops to <8%, or we conclusively show macro doesn't help and remove the unused code path. Either result is actionable.",
    ),
    dict(
        number="M8",
        title="Churn label is binary; reality is continuous",
        severity="MEDIUM",
        tags=["label design", "business intent"],
        plain="A customer is 'churned' if they have no invoice in 6 months AND no won quote in 3 months. So a €500k account that drops to €5k once a year is 'active'. So is a €500k account that drops to one €50k order. Both are very different stories from sales' perspective.",
        impact="We catch full silence well (AUC 0.94) but miss revenue erosion. A customer can lose 80% of their wallet without flipping the label. The at_risk_revenue calculation is 'P(silent) × forecast revenue' — it doesn't capture 'P(decline by 50%) × historical wallet'.",
        fix="Add a second model: predict next-12m revenue decline ratio (decline_ratio = next_12m_rev / trailing_12m_rev). Customers with decline_ratio < 0.5 are 'at-risk-of-major-decline'. This is a regression problem; same features work. Surface both signals: P(churn) and P(major decline) — they're different actions.",
    ),
    dict(
        number="M9",
        title="Churn training uses 3 rolling as-of dates 3 months apart — the same customers appear multiple times",
        severity="MEDIUM",
        tags=["leakage", "CV"],
        plain="We train on labels from 3 rolling cutoffs (e.g., Oct-2024, Jul-2024, Apr-2024). The feature vectors for the same customer at those dates are computed from overlapping trailing windows. The TimeSeriesSplit then splits this stacked frame by row index, not by customer or date.",
        impact="Cross-validation reports optimistic AUC because the test set can contain a customer that's also in the train set with a slightly different as-of. The model has 'seen' them. Real-world AUC will be lower than 0.94 / 0.81 / 0.79 — probably by 3-5 points.",
        fix="(a) Use only one as-of per customer (the most recent that has full label visibility). Costs label volume but eliminates double-counting. Or (b) custom CV splitter that holds out by customer_id, not by row. Or (c) collapse to one row per customer per as-of and use group-aware CV.",
    ),
    dict(
        number="M10",
        title="Backtest horizon is one step ahead only",
        severity="MEDIUM",
        tags=["horizon mismatch"],
        plain="Walk-forward backtests refit at each month and predict ONLY the next month. But we DELIVER a 12-month forecast. So we measure 1-month error but quote 12-month forecasts.",
        impact="MAPE you see on the report is the easy version (1-month-ahead). The actual error 12 months out is much larger. Plans built around p50 month-12 are riskier than they appear.",
        fix="Multi-horizon backtest: at each origin, predict t+1, t+3, t+6, t+12. Report MAPE separately per horizon. Users see honest 'short-term accuracy' vs 'long-term accuracy' numbers.",
    ),
    dict(
        number="M11",
        title="No model drift / decay monitoring",
        severity="LOW",
        tags=["governance", "ops"],
        plain="We train once and produce forecasts. There's no mechanism to check whether last quarter's forecasts actually matched reality, or to flag a model whose error has tripled since deployment.",
        impact="A degraded model can quietly produce wrong numbers for months before anyone notices. The business loses trust when a quarter goes badly and nobody had told them the model was off.",
        fix="Persist each forecast vintage in a small table (forecast_date, target_date, p50, p80_low, p80_high, actual_when_known). A monthly job appends actuals and computes rolling error. Threshold alert when WAPE drifts above 1.5x trained value.",
    ),
    dict(
        number="M12",
        title="Log-transform is applied unconditionally and back-transform isn't bias-corrected",
        severity="LOW",
        tags=["transform"],
        plain="For revenue / quantity we fit on log1p(value) then expm1 the forecast. Mean(log(X)) ≠ log(Mean(X)) — this is Jensen's inequality. Our back-transformed p50 is biased low because the median of the log-space distribution maps to the median of the original distribution, but the mean (which we report as p50) does not.",
        impact="Revenue forecasts are slightly conservative — typically 2-8% low on a multiplicative basis. Adds up to material euros on €1M+ accounts over a year.",
        fix="Add the standard correction: forecasted_value = exp(mu + sigma^2/2) where sigma is the in-sample residual std in log space. Documents the assumption clearly. Or, if we move to quantile regression (M2), the back-transform issue goes away because we're predicting quantiles directly.",
    ),
    dict(
        number="M13",
        title="Single global model per series; no transfer learning across similar customers",
        severity="LOW",
        tags=["pooling", "small sample"],
        plain="A €30k/month customer with 18 months of history gets its own EMA model. A €5k/month customer with 6 months gets its own EMA. They could be in the same industry, ordering similar products — but the models don't share strength.",
        impact="Small / new customers can't benefit from patterns learned across the customer base. Cold-start is poor: a customer who appeared 3 months ago gets a flat forecast.",
        fix="Hierarchical Bayesian / mixed-effects model: shared parameters across customers within a cluster (industry, size tier, commodity mix), customer-specific deviations on top. Practical implementation: lightgbm with customer_id + cluster features (already done for churn — extend to forecasting). Or use a Prophet-style model with a customer-fixed-effect.",
    ),
]


DATA_ISSUES = [
    dict(
        number="D1",
        title="Customer names anonymised — all 1,438 records have NaN name",
        severity="HIGH",
        tags=["business usability"],
        plain="Every customer record in customers.parquet has a name field set to null. We have customer_id (a number) but no human-readable name. So when the churn predictions surface 'customer 105134 at 99% churn risk' the salesperson has to look up who that is.",
        impact="Slows down activation — every list needs an offline lookup. Reduces trust ('is this customer even real?'). Breaks the at-risk-revenue narrative ('we're losing customer 105134' has zero emotional weight).",
        fix="Source — you already noted this. If anonymisation is required for the data sharing, build a separate lookup table (kept by you, not in the model) that maps customer_id → name and join it at view time. The model never sees the name.",
    ),
    dict(
        number="D2",
        title="76% of customers (735 of 967) have ≤3 invoices total",
        severity="HIGH",
        tags=["sparsity", "fundamental"],
        plain="Of the 967 customers with any invoice in 4 years, 451 have a single invoice and 735 have three or fewer. Only 40 customers have 24+ invoices. The top-50 customers account for 67.7% of revenue.",
        impact="The customer-level forecasts are statistically meaningless for ~90% of accounts. Even with the right model family (M4), a customer with 1-3 data points cannot be forecast. The methodology is mostly carrying these as 'other' buckets.",
        fix="Source — if the long tail is genuinely transactional (one-off prototypes, walk-in orders), there's nothing to fix at the data level — accept it and forecast only the top tier. If the long tail represents fragmentation across multiple business units / billing entities, link them: same parent company should aggregate. Worth a session with finance to verify.",
    ),
    dict(
        number="D3",
        title="71% of SKUs (871 of 1,221) have ≤3 invoices total",
        severity="HIGH",
        tags=["sparsity", "fundamental"],
        plain="Same story for articles: most are made-to-order or extremely rare. Only 102 SKUs have 12+ invoice lines. The top-100 captures most of the volume.",
        impact="Per-SKU forecasts for the long tail are guess-work. The proportional bucket roll-down we use is a coping mechanism, not a forecast. SKU-level decisions (which products to promote, which to discontinue) cannot be data-driven for the tail.",
        fix="Source — group SKUs into product families (already partially done via commodity_group). Bring the article master closer to a product hierarchy (family → variant → SKU) so we forecast at the family level and decompose by mix.",
    ),
    dict(
        number="D4",
        title="Only 48 months of history",
        severity="MEDIUM",
        tags=["short history"],
        plain="Invoice history runs Jan 2022 to Dec 2025. That's 48 months. For annual seasonality we need at least 24-36 months and ideally 3+ full cycles. For SARIMAX with order (1,1,1) and 2 exog at month 12+ horizon we're statistically on thin ice.",
        impact="Seasonal model components (ETS additive seasonal) can fit but with low confidence. The reported MAPEs likely underestimate real out-of-sample error because the backtest windows are short. Cross-COVID-recovery effects (2022→2023) still dominate trend estimates.",
        fix="Source — request 2019-2021 data if available. Even noisy older years would stabilise the seasonality estimates. If not available, treat all results as 'first 2 years of operation' and re-train monthly.",
    ),
    dict(
        number="D5",
        title="Quote data has no 'open' status — pipeline forecast is a proxy",
        severity="HIGH",
        tags=["pipeline accuracy"],
        plain="Quotes are tagged 'won' (1,684) or 'lost' (2,855) — every quote has been resolved. There is no live pipeline of in-flight quotes. To generate the pipeline forecast we fall back to 'last 90 days of quotes' as a proxy.",
        impact="The €450k expected booked revenue for Q3+Q4 2025 is not actually a forecast on the open pipeline — it's an estimate of what the last 90 days of resolved quotes would have produced if they'd all been pending. That number can be 2x off in either direction.",
        fix="Source — surface the open-quote status in the export. The ERP/CRM almost certainly has 'sent to customer / under review / awaiting PO' states. Without these, we cannot honestly forecast pipeline; we can only describe what win-rates have been historically.",
    ),
    dict(
        number="D6",
        title="No lost-quote reason codes (rejection_code is present but coverage is low)",
        severity="MEDIUM",
        tags=["feature engineering", "actionability"],
        plain="The quotes table has a rejection_code column. Of 2,855 lost quotes, most have either missing or unreliable rejection_code values (per dq_any_issue flag). So we don't know why we lose quotes: price too high? wrong specification? competitor preference?",
        impact="The quote conversion model can predict win/lose, but cannot say which lever moves the probability. 'Customer X has 30% win rate; should we discount?' — we have no signal on whether price or other factors drive their loss decisions.",
        fix="Source — make rejection_code mandatory in the quoting workflow, with a short controlled vocabulary (price, leadtime, specification, competitor, no-decision). After 6-12 months of clean data, add it as a feature.",
    ),
    dict(
        number="D7",
        title="No customer attributes (industry, size, geography)",
        severity="HIGH",
        tags=["segmentation"],
        plain="The customers table has 3 columns: customer_id, first_seen_date, name (all NaN). No industry. No size band. No region. No B2B relationship type (OEM / distributor / direct).",
        impact="Cannot segment churn predictions by who they are. Cannot build look-alike features ('similar customers have churned at X%'). Pricing strategy cannot vary by segment because we have no segments. Marketing can't prioritise outreach by industry.",
        fix="Source — enrich the customer master with at least: industry (NACE / SIC code), country, employee-size band, customer-type (OEM / distributor / end-user). 30-60 min per customer name if done manually; faster with a data-enrichment service. Top-50 alone is the priority.",
    ),
    dict(
        number="D8",
        title="No competitor / market-share / win-loss-to-whom data",
        severity="MEDIUM",
        tags=["external context"],
        plain="When we lose a quote we know the customer chose somebody else, but not whom. We have no panel data on competitor pricing, market-share movements, or competitor product launches.",
        impact="Pricing decisions can't be benchmarked. 'Are we under-priced? Over-priced?' — no way to answer from internal data alone.",
        fix="Source — at minimum, capture lost-to-competitor as a structured field on lost quotes. Long term, subscribe to a market-data provider (VDMA produces German industrial price indices; broader options exist).",
    ),
    dict(
        number="D9",
        title="No price-list data; we can only observe transacted prices",
        severity="MEDIUM",
        tags=["elasticity"],
        plain="We see what customers paid (revenue_per_unit). We don't see what the list price was, what was discounted, or what the discount was at quote time vs final. So we can't compute discount-driven volume sensitivity.",
        impact="The forecast can't disentangle 'we shipped 1000 pieces because we priced low' from 'we shipped 1000 pieces because demand was high'. Price-elasticity work (Studio screen) is operating without ground truth.",
        fix="Source — extract the SAP price list (likely a periodic table). For each invoice line, attach the list price valid on that date. Discount = (list - transacted) / list. After 6 months of this attached, elasticity models become feasible.",
    ),
    dict(
        number="D10",
        title="No order-level fulfillment data (lead time, delivery vs request date)",
        severity="LOW",
        tags=["operational"],
        plain="We have invoice date (when billed) but not order date, requested delivery date, actual delivery date, or fulfillment quantity vs ordered quantity.",
        impact="Cannot diagnose churn driven by operational issues (late delivery → customer leaves). Cannot forecast revenue at the order level (earlier in the customer journey).",
        fix="Source — request the order header table from SAP. Should be 1-2 days of IT effort. Once available, add 'avg_leadtime_days' and 'on_time_delivery_rate' as customer features in churn model — high-value signal.",
    ),
    dict(
        number="D11",
        title="3 invoice rows have NaN commodity_group; 53 rows flagged dq_any_issue",
        severity="LOW",
        tags=["data quality"],
        plain="Tiny data-quality holes: 3 invoices not classified into a commodity group, 53 with at least one DQ flag (missing margin / negative margin / suspiciously low margin). We exclude these from forecasting via dq_any_issue but don't audit why they exist.",
        impact="Negligible volume impact (~1% of rows). Hides upstream system noise — there may be a category of order that systematically misses commodity classification (e.g., spare parts, sample orders).",
        fix="Source — review the 53 flagged rows manually; categorise them ('legitimate concession', 'data entry error', 'spare-parts category'). Update the ETL to handle each category. Document for audit.",
    ),
    dict(
        number="D12",
        title="No working-day / calendar / plant-shutdown feature",
        severity="LOW",
        tags=["seasonality"],
        plain="A month with 18 working days is materially different from a month with 23 working days for industrial output. We have no calendar feature, no holiday flags, no plant-vacation indicator.",
        impact="Monthly forecasts conflate working-day variation with demand variation. August (vacation) and December (Christmas shutdown) get treated as 'low months' by the seasonal model — but the underlying customer demand may be unchanged, just compressed.",
        fix="Source — easy. Add a fixed lookup: month → working_days_DE. Also flag known shutdown periods. Becomes an exog regressor for the forecast.",
    ),
    dict(
        number="D13",
        title="No PO / contract length data",
        severity="MEDIUM",
        tags=["churn signal"],
        plain="We see invoices and quotes. We don't see whether a customer is on a multi-year supply contract, a frame agreement, or quoting case-by-case. A customer on a 2-year contract has structurally different churn dynamics from a spot buyer.",
        impact="Churn model treats both equally. A spot buyer with a 6-month silence may be normal; a contract customer with a 6-month silence may already have churned. We're under-weighting the relationship type.",
        fix="Source — request the contract / frame-agreement master. Add 'on_active_contract' and 'contract_end_date' as customer features. Should noticeably lift churn AUC and (more importantly) reduce false positives on contract customers.",
    ),
]


CROSS_CUTTING = [
    ("Model governance",
     "Right now models are trained ad-hoc when the notebook runs. There's no model registry, no versioning, no record of 'which model produced which forecast'. If forecasts feed into a business decision and the decision goes badly, we can't audit which model was responsible. Pryzm should adopt MLflow / a simple JSON manifest per run."),
    ("Reproducibility",
     "The notebook is reproducible only if all data is on disk and Python packages are pinned. There's no notebook test (CI) that exercises the full pipeline. A regression in forecasting_extensions.py would only surface when somebody re-runs the notebook."),
    ("Confidence calibration audit",
     "We say p80 means '80% of actuals will fall in the band'. We never check that this is true. After 12 months of forecasts vs actuals we should compute the empirical coverage — if only 60% of actuals fell in the p80 band, the bands are wrong."),
    ("Action loop closure",
     "Churn predictions surface customers at risk. We don't track what happens next. Did the salesperson contact them? Did the customer come back? Without this loop we can't measure the model's business impact (did it save revenue?) or improve it (treatment-effect data)."),
]


def build() -> None:
    doc = Document()

    title = doc.add_heading("Pryzm Forecast & Churn — Methodology Review", 0)
    for r in title.runs:
        r.font.color.rgb = ROSE
    para(doc, f"Generated {datetime.now(timezone.utc):%Y-%m-%d}. Independent self-critique of what we built. Every issue is tagged METHODOLOGY (something we can fix in code) or DATA (something missing from the source). Severity is calibrated against business impact, not academic purity.",
         italic=True, color=MUTED)

    heading(doc, "Executive summary", level=1)
    para(doc,
         "We shipped a working forecast + churn + pipeline system. It passes the primary accuracy gate, churn AUC is 0.79-0.94, quote conversion AUC is 0.84. "
         "The headline numbers are usable for decision-making AT THE RIGHT LEVEL.")
    para(doc,
         "But we have 13 methodology issues and 13 data gaps. They cluster into four themes:",
         bold=True)
    bullets(doc, [
        "Calibration is optimistic. AUCs and MAPEs we report are best-case (leakage in M1, M9; selection bias in M3; Normal-band assumption in M2).",
        "Per-customer / per-SKU work is hobbled by the wrong model family for bursty data (M4). Fix the model, the 'informational tier' becomes usable.",
        "The customer base is genuinely sparse (D2, D3): 76% of customers have ≤3 invoices. Some forecasts are impossible regardless of methodology.",
        "Key context is missing from the source (D1, D5, D7, D13): customer names, open-quote status, customer attributes, contract relationships. Each unlocks meaningful lift.",
    ])
    para(doc, "Prioritised fix list (top-5)", bold=True)
    bullets(doc, [
        "1. Fix quote-conversion leakage (M1). 1 day. Will reveal true AUC.",
        "2. Add Croston/TSB for bursty per-customer/SKU forecasts (M4). 2 days. Big lift on per-customer at-risk-revenue.",
        "3. Source customer names and at least industry attribute (D1, D7 subset). Business effort, not engineering.",
        "4. Source open-quote status from ERP (D5). Days, not weeks. Replaces proxy pipeline with real pipeline.",
        "5. Bootstrap quantile bands instead of Normal-approx (M2). 1 day. Honest uncertainty bands.",
    ])

    doc.add_page_break()
    heading(doc, "Methodology issues (things we can fix)", level=1)
    para(doc, "13 issues across model design, validation, calibration, and governance. Sorted by severity.", italic=True, color=MUTED)
    for it in sorted(METHODOLOGY_ISSUES, key=lambda x: ({"HIGH":0,"MEDIUM":1,"LOW":2}[x["severity"]], x["number"])):
        issue(doc, **it)

    doc.add_page_break()
    heading(doc, "Data gaps (things missing from the source)", level=1)
    para(doc, "13 issues. Some are easy to fix (working days, rejection codes); some require organisational work (customer enrichment, contract data); some are structural realities (long-tail sparsity).", italic=True, color=MUTED)
    for it in sorted(DATA_ISSUES, key=lambda x: ({"HIGH":0,"MEDIUM":1,"LOW":2}[x["severity"]], x["number"])):
        issue(doc, **it)

    doc.add_page_break()
    heading(doc, "Cross-cutting concerns", level=1)
    for name, body in CROSS_CUTTING:
        heading(doc, name, level=2)
        para(doc, body)

    heading(doc, "Priority matrix", level=1)
    para(doc, "Quadrant view: which gaps move the needle most for the effort they cost.", italic=True, color=MUTED)
    matrix = [
        ("HIGH impact / LOW effort",
         ["M1 quote-conversion leakage fix",
          "M2 bootstrap quantile bands",
          "D12 add working-day calendar feature",
          "D11 audit the 53 DQ-flagged rows"]),
        ("HIGH impact / MEDIUM effort",
         ["M4 Croston/TSB for intermittent series",
          "M8 add revenue-decline label alongside churn",
          "M9 fix churn CV with customer-grouped splits",
          "D6 add lost-quote reason codes (workflow change)",
          "D10 add order-level fulfillment data"]),
    ]
    matrix2 = [
        ("HIGH impact / HIGH effort",
         ["D1 + D7 enrich customer master (names + industry + size + region)",
          "D5 source open-quote pipeline data from ERP",
          "D9 source price-list data and compute discount features",
          "D13 source contract / frame-agreement master",
          "M13 hierarchical / pooled forecasting"]),
        ("LOW-MEDIUM impact",
         ["M3 nested selection split",
          "M5 MinT reconciliation",
          "M6 quarterly multivariate sampling",
          "M7 SARIMAX diagnostic for BKAIZ",
          "M10 multi-horizon backtest",
          "M11 model drift monitoring",
          "M12 log-transform bias correction",
          "D2 D3 long-tail sparsity (mostly inherent)",
          "D4 longer history (request 2019-2021)",
          "D8 competitor / market-share data"]),
    ]
    for label, items in matrix + matrix2:
        heading(doc, label, level=2)
        bullets(doc, items)

    doc.add_page_break()
    heading(doc, "What this review is NOT", level=1)
    bullets(doc, [
        "Not a criticism of the data being 'bad'. The data is what the business records. We're flagging what would lift performance — for the business to weigh against the cost of collection.",
        "Not a recommendation to rebuild from scratch. The system works. Every issue here is incremental on top of what shipped.",
        "Not a list of things blocking go-live. The primary gate passes. Churn and conversion models are usable today. The fixes are about going from 'usable' to 'auditable enterprise-grade'.",
        "Not exhaustive. There are smaller issues (no SHAP / no explanations on individual churn predictions, no A/B framework, no canary on pricing recommendations). Surfacing those would multiply the page count without changing the picture.",
    ])

    heading(doc, "Glossary", level=1)
    bullets(doc, [
        "AUC: how well a yes/no model separates the two classes. 0.5 = random, 1.0 = perfect. Above 0.70 useful, above 0.80 strong.",
        "MAPE: average percentage error. Good for ratios. Bad for bursty positive numbers.",
        "WAPE: total error divided by total actual. Robust to bursty positive numbers.",
        "p50/p80/p95: forecast quantiles (most-likely / 80% range / 95% range).",
        "Croston / TSB: forecasting methods built for intermittent demand (lots of zero-months interrupted by occasional orders).",
        "MinT: minimum-trace reconciliation. A way to make sub-totals add up to totals without throwing away local signal.",
        "SARIMAX: seasonal ARIMA with exogenous regressors. A classical time-series model.",
        "Bootstrap: estimate uncertainty by resampling. No distribution assumption needed.",
        "Feature leakage: training feature accidentally contains information from the target (the thing you're trying to predict). Makes the model look better than it really is.",
        "Walk-forward backtest: validate a forecast by pretending you don't know the last N months, predicting them, and comparing. Repeated through history.",
        "Survival analysis: models how long until an event happens (here: until a customer churns). More nuanced than fixed-horizon classification.",
    ])

    doc.save(str(REPORT))
    print(f"Wrote {REPORT}  ({REPORT.stat().st_size/1024:.1f} KB)")


if __name__ == "__main__":
    build()
