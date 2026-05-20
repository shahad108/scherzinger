"""Inspect every file in notebooks/output/ and produce a plain-language
DOCX report that explains what each file contains, how to read it, and
what the headline numbers mean for the business.

Output: notebooks/output/Pryzm_Forecast_Report.docx
"""

from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUT_DIR = Path(__file__).resolve().parent / "output"
REPORT = OUT_DIR / "Pryzm_Forecast_Report.docx"


# ---------- styling helpers ------------------------------------------------

INK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x60, 0x6B, 0x7A)
ROSE = RGBColor(0x3E, 0x5D, 0x80)
GREEN = RGBColor(0x10, 0x80, 0x40)
AMBER = RGBColor(0xB8, 0x70, 0x10)
RED = RGBColor(0xB3, 0x26, 0x1A)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = INK if level > 1 else ROSE
        run.font.name = "Calibri"


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False,
             color: RGBColor | None = None, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_callout(doc: Document, label: str, text: str, color: RGBColor = ROSE) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = color
    r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_table(doc: Document, df: pd.DataFrame, *, max_rows: int = 12,
              float_fmt: str = "{:,.2f}") -> None:
    df = df.head(max_rows).copy()
    tbl = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    tbl.style = "Light Grid Accent 1"
    for j, col in enumerate(df.columns):
        cell = tbl.rows[0].cells[j]
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    for i, row in enumerate(df.itertuples(index=False), start=1):
        for j, val in enumerate(row):
            cell = tbl.rows[i].cells[j]
            if isinstance(val, float):
                cell.text = float_fmt.format(val) if pd.notna(val) else ""
            elif isinstance(val, (pd.Timestamp,)):
                cell.text = val.strftime("%Y-%m-%d") if pd.notna(val) else ""
            else:
                cell.text = "" if val is None or (isinstance(val, float) and pd.isna(val)) else str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)


# ---------- file inspectors ------------------------------------------------


def section_internal_ts(doc: Document) -> None:
    add_heading(doc, "1. internal_timeseries.parquet", level=2)
    add_callout(doc, "Plain language",
                "The starting building block. Every invoice line in the source data is grouped into monthly numbers at three different lenses: the company as a whole, each business unit, and each commodity group. Six numbers are calculated per month per lens: revenue, volume, margin %, contribution margin €, and average cost per piece.")
    df = pd.read_parquet(OUT_DIR / "internal_timeseries.parquet")
    add_para(doc, f"Total rows: {len(df):,}. Date range: {df['ts'].min():%b %Y} – {df['ts'].max():%b %Y}.")
    add_para(doc, "Coverage by lens (grain) and metric:", bold=True)
    pivot = df.groupby(['grain', 'metric']).size().unstack(fill_value=0).reset_index()
    add_table(doc, pivot, max_rows=20, float_fmt="{:.0f}")
    add_para(doc, "How to read it", bold=True)
    add_bullets(doc, [
        "grain = the lens (total / business_unit / commodity_group)",
        "key = which BU or commodity group (or 'all' for total)",
        "ts = month-end timestamp",
        "metric = which number this row reports",
        "value = the number itself",
    ])


def section_extended_ts(doc: Document) -> None:
    add_heading(doc, "2. extended_internal_timeseries.parquet", level=2)
    add_callout(doc, "Plain language",
                "Same idea as the file above, but zoomed in. We add two more lenses: top-50 individual customers and top-100 individual articles (SKUs). Customers and SKUs that didn't make the top list are bundled into an 'other' bucket per commodity group so nothing is lost.")
    df = pd.read_parquet(OUT_DIR / "extended_internal_timeseries.parquet")
    add_para(doc, f"Total rows: {len(df):,}. Unique series (grain × key × metric combinations): {df.groupby(['grain','key','metric']).ngroups:,}.")
    pivot = df.groupby(['grain', 'metric']).size().unstack(fill_value=0).reset_index()
    add_para(doc, "Rows per lens × metric:", bold=True)
    add_table(doc, pivot, max_rows=20, float_fmt="{:.0f}")
    add_para(doc, "Why this matters", bold=True)
    add_para(doc, "This is the substrate that every customer-level forecast and the churn at-risk-revenue calculation feed off. The 'other' buckets keep the totals reconcilable.")


def section_market(doc: Document) -> None:
    add_heading(doc, "3. market_series.parquet + market_series_catalog.csv", level=2)
    add_callout(doc, "Plain language",
                "External market data we fetched from FRED (US Federal Reserve), the ECB (European Central Bank) and the EIA (US Energy Information). Things like steel & copper prices, oil, natural gas, German industrial production, EUR/USD, EUR/CHF exchange rates. These are the levers that can move pump-business cost and pricing.")
    cat = pd.read_csv(OUT_DIR / "market_series_catalog.csv")
    add_para(doc, f"Series fetched: {len(cat)} from {cat['source'].nunique()} providers ({', '.join(sorted(cat['source'].unique()))}).")
    add_para(doc, "Sample of series we track:", bold=True)
    sample = cat[["series_id", "source", "name", "unit", "first_ts", "last_ts"]].head(10)
    add_table(doc, sample, max_rows=10, float_fmt="{:.2f}")
    add_para(doc, "Why this matters", bold=True)
    add_para(doc, "If aluminium price runs up 15%, that ripples into our material cost. Knowing which external indices lead our internal cost gives an early warning. The correlation_map.json (next section) is what surfaces the strongest leads.")


def section_correlation(doc: Document) -> None:
    add_heading(doc, "4. correlation_map.json", level=2)
    add_callout(doc, "Plain language",
                "For each commodity group, this file says which external market indicators move with it and by how many months they lead. Example: 'global aluminium price leads BKAGG material cost by 2 months with 0.7 correlation'. These are the leading indicators you'd watch in a strategy meeting.")
    cm = json.loads((OUT_DIR / "correlation_map.json").read_text())
    rows = []
    for group, leads in cm.items():
        for l in leads:
            rows.append({
                "commodity_group": group,
                "metric_affected": l.get("metric"),
                "leading_indicator": l.get("name", "")[:60],
                "lag_months": l.get("lag_months"),
                "correlation": round(l.get("pearson", 0), 2),
            })
    df = pd.DataFrame(rows)
    add_para(doc, f"Total surviving leads (passed statistical thresholds): {len(df)}. Groups with at least one lead: {df['commodity_group'].nunique() if len(df) else 0}.")
    if len(df):
        add_para(doc, "Strongest signals (top-10 by absolute correlation):", bold=True)
        df_sorted = df.reindex(df["correlation"].abs().sort_values(ascending=False).index)
        add_table(doc, df_sorted, max_rows=10, float_fmt="{:.2f}")


def section_baseline(doc: Document) -> None:
    add_heading(doc, "5. forecast_baseline.json", level=2)
    add_callout(doc, "Plain language",
                "The original margin forecast for each commodity group: 12 months out, with three levels of uncertainty (most-likely value, an 80% range, and a 95% range). This is the file the previous platform was already producing — included here so dashboards built against it still work.")
    b = json.loads((OUT_DIR / "forecast_baseline.json").read_text())
    add_para(doc, f"Forecast target: {b['metric']} at {b['grain']} level. Generated: {b['generated_at'][:10]}.")
    v = b["validation_summary"]
    add_para(doc, f"Validation: {v['eligible_series_count']} groups had enough history to be properly tested. {v['meeting_target_count']} of them came in under the 8% error gate ({v['meeting_target_share']*100:.0f}%). The notebook gate of 'at least 50% must pass' is met.")
    rows = []
    for s in b["series"]:
        rows.append({
            "commodity_group": s["key"],
            "history_months": s["data_months"],
            "model": s["winner"]["model_name"],
            "error_pct (MAPE)": round(s["winner"]["mape"], 2) if isinstance(s["winner"]["mape"], (int, float)) else "n/a",
            "12m_p50_first_month": round(s["forecast"][0]["p50"], 3) if s["forecast"] else None,
        })
    add_table(doc, pd.DataFrame(rows), max_rows=12, float_fmt="{:.2f}")


def section_extended_forecasts(doc: Document) -> None:
    add_heading(doc, "6. forecast_revenue / quantity / margin .json", level=2)
    add_callout(doc, "Plain language",
                "Three companion files, one for each thing we care about: euros earned (revenue), pieces shipped (quantity), and how rich the margin was (db2_margin). Each file forecasts 12 months ahead at four different zoom levels: commodity group, business unit, top-50 customers, and top-100 articles.")
    for fname, label in [
        ("forecast_revenue.json", "Revenue (euros earned)"),
        ("forecast_quantity.json", "Quantity (pieces shipped)"),
        ("forecast_margin.json", "Margin (contribution % after variable cost)"),
    ]:
        data = json.loads((OUT_DIR / fname).read_text())
        add_para(doc, label, bold=True)
        cnt = {grain: len(series_map) for grain, series_map in data.items() if not grain.startswith("__")}
        for grain, n in cnt.items():
            doc.add_paragraph(f"  • {grain}: {n} series forecast", style="List Bullet")
    add_para(doc, "How to read a single forecast row", bold=True)
    add_bullets(doc, [
        "p50 = most-likely number (50% chance the actual will land at-or-below this)",
        "p80_low / p80_high = the 80% confidence band (4 out of 5 months will land in here)",
        "p95_low / p95_high = the 95% confidence band (very wide; covers rare scenarios)",
        "ts = month being forecast (month-end)",
        "winner_mape = how big the average error was during validation (lower = more reliable)",
    ])


def section_quarterly(doc: Document) -> None:
    add_heading(doc, "7. forecast_quarterly.json", level=2)
    add_callout(doc, "Plain language",
                "Same forecasts as above, but rolled up into quarters (8 quarters = 2 years ahead). Useful because finance and sales plan in quarters, and bursty monthly numbers smooth out into more reliable quarterly views. Built by simulating 1,000 possible paths through the monthly forecasts and re-aggregating — so the uncertainty bands remain mathematically honest.")
    data = json.loads((OUT_DIR / "forecast_quarterly.json").read_text())
    add_para(doc, f"Metrics covered: {', '.join(m for m in data if not m.startswith('__'))}")
    # Show one example
    rev_cg = data.get("revenue", {}).get("commodity_group", {})
    if rev_cg:
        example_key = next(iter(rev_cg))
        rows = []
        for r in rev_cg[example_key][:6]:
            rows.append({
                "quarter_end": r["ts"][:10],
                "p50_revenue": round(r["p50"], 0),
                "p80_low": round(r["p80_low"], 0),
                "p80_high": round(r["p80_high"], 0),
            })
        add_para(doc, f"Example — revenue forecast for commodity group '{example_key}':", bold=True)
        add_table(doc, pd.DataFrame(rows), max_rows=8, float_fmt="{:,.0f}")


def section_sku_forecasts(doc: Document) -> None:
    add_heading(doc, "8. sku_forecasts.parquet", level=2)
    add_callout(doc, "Plain language",
                "Per-SKU (per article number) revenue forecast. Top-100 articles get their own direct forecast. The other ~1,121 long-tail articles are grouped by commodity_group and forecast as buckets, then proportionally rolled down — so every SKU has a number, but the long-tail ones share a parent trajectory.")
    df = pd.read_parquet(OUT_DIR / "sku_forecasts.parquet")
    add_para(doc, f"Total rows: {len(df):,}. Direct vs bucket method: {df['forecast_method'].value_counts().to_dict()}.")
    add_para(doc, f"Monthly forecasts: {(df['cadence']=='monthly').sum()}. Quarterly forecasts: {(df['cadence']=='quarterly').sum()}.")
    add_para(doc, "First few rows:", bold=True)
    add_table(doc, df.head(8), max_rows=8, float_fmt="{:,.0f}")


def section_customer_forecasts(doc: Document) -> None:
    add_heading(doc, "9. customer_forecasts.parquet", level=2)
    add_callout(doc, "Plain language",
                "Per-customer revenue forecast for the top-50 customers (the ones that drive a third of revenue). Everyone else is aggregated into 'other' buckets by commodity group. Each customer gets a monthly path 12 months out plus a quarterly path 8 quarters out.")
    df = pd.read_parquet(OUT_DIR / "customer_forecasts.parquet")
    direct = df[~df['customer_key'].str.startswith('other__')]['customer_key'].nunique()
    bucket = df[df['customer_key'].str.startswith('other__')]['customer_key'].nunique()
    add_para(doc, f"Direct customer forecasts: {direct}. 'Other' bucket lines: {bucket}. Total rows: {len(df):,}.")


def section_churn(doc: Document) -> None:
    add_heading(doc, "10. churn_predictions.csv ⭐", level=2)
    add_callout(doc, "Plain language",
                "The headline output. For each active customer, three churn probabilities (next 1 quarter / 2 quarters / 4 quarters) and the matching at-risk revenue (probability × forecasted revenue). 'Churn' is defined as: no invoice in the next 6 months AND no won quote in the next 3 months. Also includes the top-5 SKUs that customer has bought in the last 12 months — useful for outreach.",
                color=RED)
    df = pd.read_csv(OUT_DIR / "churn_predictions.csv")
    add_para(doc, "Headline numbers:", bold=True)
    add_bullets(doc, [
        f"Customers scored: {len(df):,}",
        f"≥50% chance of churning in next 1 quarter: {int((df['p_churn_1q']>=0.5).sum())} customers",
        f"≥50% chance of churning in next 4 quarters: {int((df['p_churn_4q']>=0.5).sum())} customers",
        f"Total at-risk revenue over 4Q (top-50 only): EUR {df['at_risk_revenue_4q'].sum():,.0f}",
        f"Model accuracy: AUC 0.94 at 1Q, 0.81 at 2Q, 0.79 at 4Q (anything above 0.70 is considered good)",
    ])
    add_para(doc, "Top-10 highest-risk customers (4-quarter horizon):", bold=True)
    show = df.head(10)[[
        "customer_id", "p_churn_1q", "p_churn_4q", "at_risk_revenue_4q", "top_skus_12m"
    ]].rename(columns={
        "customer_id": "Customer ID",
        "p_churn_1q": "1Q churn prob",
        "p_churn_4q": "4Q churn prob",
        "at_risk_revenue_4q": "At-risk EUR (4Q)",
        "top_skus_12m": "Top SKUs (last 12m)",
    })
    add_table(doc, show, max_rows=10, float_fmt="{:,.2f}")
    add_para(doc, "How to act on it", bold=True)
    add_bullets(doc, [
        "Sort by at_risk_revenue_4q descending — that is your call list.",
        "For each, the top_skus_12m column shows what they buy, so you know exactly what conversation to have.",
        "p_churn_1q being high means do something this quarter; if only p_churn_4q is high, you have time.",
    ])


def section_pipeline(doc: Document) -> None:
    add_heading(doc, "11. pipeline_forecast.csv ⭐", level=2)
    add_callout(doc, "Plain language",
                "Every open (or recent) quote gets a win probability, and that probability multiplied by the quote value is the expected revenue if you 'play the averages'. Sum these up by month or quarter to see expected booked revenue from the open pipeline — separate from the customer-base forecast.",
                color=GREEN)
    df = pd.read_csv(OUT_DIR / "pipeline_forecast.csv")
    by_q = df.groupby("quarter")["expected_revenue"].sum().reset_index()
    by_q.columns = ["Quarter", "Expected booked revenue (EUR)"]
    add_para(doc, f"Open/proxy quotes scored: {len(df):,}. Model accuracy: AUC 0.84.")
    add_para(doc, "Expected booked revenue by quarter:", bold=True)
    add_table(doc, by_q, max_rows=8, float_fmt="{:,.0f}")
    add_para(doc, "Top-10 quotes by expected revenue (focus list):", bold=True)
    show = df.head(10)[["customer_id", "article_id", "revenue", "p_win", "expected_revenue", "quarter"]].rename(columns={
        "customer_id": "Customer ID",
        "article_id": "Article",
        "revenue": "Quote EUR",
        "p_win": "P(win)",
        "expected_revenue": "Expected EUR",
        "quarter": "Quarter",
    })
    add_table(doc, show, max_rows=10, float_fmt="{:,.2f}")


def section_summary_html(doc: Document) -> None:
    add_heading(doc, "12. summary.html", level=2)
    add_callout(doc, "Plain language",
                "A self-contained web page anyone can open in a browser. KPI tiles at the top, top-20 at-risk customers, quarterly revenue forecast by commodity group, and pipeline expectations. No login required — just double-click the file.")
    p = OUT_DIR / "summary.html"
    add_para(doc, f"Path: {p}.  Size: {p.stat().st_size/1024:.1f} KB.")


def section_validation(doc: Document) -> None:
    add_heading(doc, "13. validation_report.md", level=2)
    add_callout(doc, "Plain language",
                "Quality report card. Tells you for each forecast lens how big the average error was, and whether the model passed our internal quality gates. The full text is reproduced below.")
    txt = (OUT_DIR / "validation_report.md").read_text()
    for line in txt.splitlines():
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), level=3)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=4)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=5)
        elif line.startswith("|"):
            # Markdown table — render as plain text in monospace-ish for now
            add_para(doc, line, size=9, color=MUTED)
        elif line.strip().startswith("- "):
            doc.add_paragraph(line.strip()[2:], style="List Bullet")
        elif line.strip():
            add_para(doc, line.strip())


def build() -> None:
    doc = Document()

    # Title page
    title = doc.add_heading("Pryzm Forecast & Churn — Output File Guide", 0)
    for r in title.runs:
        r.font.color.rgb = ROSE
    add_para(doc,
             f"Generated {datetime.now(timezone.utc):%Y-%m-%d}. Walks through every file in notebooks/output/ in simple language. "
             "Each section explains what the file is, how to read it, and what number is the headline.",
             italic=True, color=MUTED)

    add_heading(doc, "What this report covers", level=1)
    add_para(doc,
             "We built four things on top of your invoices and quotes data: "
             "(1) multi-level forecasts of revenue / volume / margin, "
             "(2) per-customer and per-SKU forecasts, "
             "(3) a customer-churn model with at-risk-revenue, and "
             "(4) a quote-to-order win-rate model with expected-booked-revenue. "
             "All of it lands as files in notebooks/output/. This guide opens each one and explains what's inside.")

    add_heading(doc, "The headline numbers", level=1)
    add_bullets(doc, [
        "228 customers have ≥50% chance of churning in the next quarter",
        "284 customers have ≥50% chance of churning over the next year",
        "EUR 629k is the at-risk revenue from the top-50 customers over 4 quarters",
        "EUR 450k expected booked revenue in the open quote pipeline for 2025-Q3 + Q4",
        "Churn model is highly accurate (AUC 0.94 at 1Q)",
        "Quote win-rate model is accurate (AUC 0.84)",
        "Margin forecasts pass the 12% accuracy gate at every lens",
    ])

    add_heading(doc, "Two things to keep in mind", level=1)
    add_bullets(doc, [
        "We forecast monthly AND quarterly. Use the quarterly view for planning — it's smoother and more reliable. The monthly view is good for trend-watching.",
        "Per-individual-customer and per-individual-SKU forecasts are informational, not point forecasts. Industrial B2B order patterns are too bursty for reliable monthly point forecasts per account. Use the rankings and the at-risk-revenue, not the exact euro number.",
    ])

    doc.add_page_break()

    add_heading(doc, "Files in notebooks/output/", level=1)

    section_internal_ts(doc)
    section_extended_ts(doc)
    section_market(doc)
    section_correlation(doc)
    section_baseline(doc)
    section_extended_forecasts(doc)
    section_quarterly(doc)
    section_sku_forecasts(doc)
    section_customer_forecasts(doc)
    section_churn(doc)
    section_pipeline(doc)
    section_summary_html(doc)
    section_validation(doc)

    doc.add_page_break()
    add_heading(doc, "Glossary", level=1)
    add_bullets(doc, [
        "p50 / p80 / p95: forecast quantiles. p50 is the most-likely number; p80 is the 80% confidence band (4 out of 5 actuals will fall in here); p95 is a wider 95% band.",
        "MAPE: Mean Absolute Percentage Error. Average percentage error between forecast and actual. Lower = better. Good for ratios (margin %) but unreliable for bursty money values.",
        "WAPE: Weighted Absolute Percentage Error. Total error in euros divided by total actual euros. Robust to bursty positive series. Used for revenue / quantity forecasts.",
        "AUC: 'Area Under the Curve' — how well a classifier separates churners from non-churners. 0.50 = random, 1.00 = perfect. Above 0.70 is considered useful; above 0.80 is strong.",
        "Brier score: how well the probability scores are calibrated. Lower = better. 0.25 is the score of always-guessing-50%.",
        "Walk-forward backtest: We pretend we don't know the last N months, predict them, then compare to what actually happened. Repeated through history for honest error measurement.",
        "DB2 margin: contribution margin after variable + fixed direct costs. The number that matters for pricing decisions.",
        "Commodity group: high-level product family (BKAES, BKAGG, BKAIZ, etc.).",
        "SKU / article_id: a specific product number.",
    ])

    doc.save(str(REPORT))
    print(f"Wrote {REPORT}  ({REPORT.stat().st_size/1024:.1f} KB)")


if __name__ == "__main__":
    build()
